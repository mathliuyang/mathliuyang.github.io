from __future__ import annotations

import argparse
import importlib.util
import re
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def load_builder(path: Path):
    spec = importlib.util.spec_from_file_location("review_builder", path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("source_docx", type=Path)
    ap.add_argument("final_docx", type=Path)
    ap.add_argument("builder_py", type=Path)
    args = ap.parse_args()

    builder = load_builder(args.builder_py)
    manuscript, title, comments = builder.read_source(args.source_docx)
    doc = Document(args.final_docx)
    failures = []

    if len(doc.tables) != 42:
        failures.append(f"Expected 42 main-story tables, found {len(doc.tables)}")
    if len(doc.sections) != 1:
        failures.append(f"Expected 1 section, found {len(doc.sections)}")

    ordered_keys = [f"R1.{i}" for i in range(1, 12)] + [f"R2.{i}" for i in range(1, 10)]
    for idx, key in enumerate(ordered_keys):
        en_table = doc.tables[2 + idx * 2]
        zh_table = doc.tables[3 + idx * 2]
        if norm(en_table.cell(0, 0).text) != "English original":
            failures.append(f"{key} English label mismatch")
        if norm(zh_table.cell(0, 0).text) != "中文译文":
            failures.append(f"{key} Chinese label mismatch")
        en_actual = [norm(p.text) for p in en_table.cell(1, 0).paragraphs if norm(p.text)]
        zh_actual = [norm(p.text) for p in zh_table.cell(1, 0).paragraphs if norm(p.text)]
        en_expected = [norm(x) for x in comments[key]]
        zh_source = builder.TRANSLATIONS[key]
        zh_expected = [norm(x) for x in (zh_source if isinstance(zh_source, list) else [zh_source])]
        if en_actual != en_expected:
            failures.append(f"{key} English mismatch: expected {en_expected!r}, got {en_actual!r}")
        if zh_actual != zh_expected:
            failures.append(f"{key} Chinese mismatch: expected {zh_expected!r}, got {zh_actual!r}")

    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        all_text += "\n" + "\n".join(cell.text for row in table.rows for cell in row.cells)
    forbidden = ["Response:", ">> We thank", "The corresponding change appears"]
    for phrase in forbidden:
        if phrase in all_text:
            failures.append(f"Forbidden author-response phrase found: {phrase}")
    if manuscript not in all_text or title not in all_text:
        failures.append("Manuscript metadata missing")

    section = doc.sections[0]
    page_cm = (section.page_width.cm, section.page_height.cm)
    margins_cm = (section.top_margin.cm, section.bottom_margin.cm, section.left_margin.cm, section.right_margin.cm)
    if any(abs(a - b) > 0.03 for a, b in zip(page_cm, (21.0, 29.7))):
        failures.append(f"Page size mismatch: {page_cm}")
    if any(abs(a - b) > 0.03 for a, b in zip(margins_cm, (2.25, 2.25, 2.35, 1.85))):
        failures.append(f"Margins mismatch: {margins_cm}")

    with zipfile.ZipFile(args.final_docx) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        settings_xml = etree.fromstring(zf.read("word/settings.xml"))
        names = set(zf.namelist())
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        if " TOC " not in document_xml or " PAGE " not in "".join(
            zf.read(n).decode("utf-8") for n in names if n.startswith("word/footer") and n.endswith(".xml")
        ):
            failures.append("Required TOC/PAGE fields missing")
        if settings_xml.find(f"{{{W}}}mirrorMargins") is None:
            failures.append("mirrorMargins missing")
        update = settings_xml.find(f"{{{W}}}updateFields")
        if update is None or update.get(f"{{{W}}}val") not in {"false", "0"}:
            failures.append("updateFields is not finalized to false")
        if not any(n.startswith("word/header") for n in names) or not any(n.startswith("word/footer") for n in names):
            failures.append("Header/footer package parts missing")

    if failures:
        print("FAIL")
        for item in failures:
            print(f"- {item}")
        raise SystemExit(1)

    print("PASS")
    print(f"Manuscript: {manuscript}")
    print(f"Reviewer comments: {len(ordered_keys)}")
    print("English bodies: 20/20 nonempty and matched")
    print("Chinese bodies: 20/20 nonempty and matched")
    print("Author responses excluded: yes")
    print(f"Main-story tables: {len(doc.tables)}")
    print(f"Page geometry (cm): {page_cm}; margins: {margins_cm}")


if __name__ == "__main__":
    main()
