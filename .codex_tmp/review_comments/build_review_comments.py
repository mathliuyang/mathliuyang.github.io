from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree

ACCENT = "1E466E"
EN_BACK = "F6F9FC"
ZH_BACK = "FAFAFA"
RULE_GRAY = "787878"
WHITE = "FFFFFF"
BLACK = "000000"
TABLE_WIDTH_DXA = 9524


TRANSLATIONS = {
    "R1.1": "第 3 页，从上往下第 11 行：对于超奇异积分方程，“α≥2”似乎应为“α > 1”。否则，请讨论缺失的情形 1 < α < 2。",
    "R1.2": "第 4 页，从上往下第 16 行，表述“G(x,t) 和 φ(t) 均连续，因此被积函数 [G(x,t)/(x−t)^α]φ(t) 也是连续函数”似乎有误。例如，若 G(x,t)=(x+t)^α 且 φ(t)=1，则 [G(x,t)/(x−t)^α]φ(t) 如何会是连续函数？",
    "R1.3": "第 4 页，从上往下第 11、12 行，区域 Ω 和 [a,x] 都是集合。那么写 Ω∈[a,x] 的目的是什么？应使用“=”或“⊆”。类似地，Ω∈[a,x−ε] 等表达也应如此处理。",
    "R1.4": "引言部分对研究动机的阐述不够充分。应明确说明所提出技术具有哪些优势。",
    "R1.5": "第 3 页，从上往下第 11 行，G(x,t)/(x−t)φ(t) 应写成 [G(x,t)/(x−t)]φ(t)，以避免混淆。文中许多地方存在类似的不准确表达。对于 G(x,t)/(x−t)^α 也请作同样修正。",
    "R1.6": "第 2 页下半部分的行距明显大于上半部分。请修正这一格式不一致问题，并确保全文行距统一。",
    "R1.7": "关于公式（9），R 是什么？它是否表示实数集？如果是，请记作 ℝ。",
    "R1.8": "作者应在数值算例部分开头说明使用了何种软件包来获得所需结果。",
    "R1.9": "作者应更详细地说明数值算例部分各数值表格和图中观察到的现象。",
    "R1.10": "所有参考文献应采用一致的格式。例如，参考文献 4 和 16 在页码前写有“pp.”，而其他参考文献没有。",
    "R1.11": "此外，稿件中仍有一些细微的语法不一致问题，需要予以修正。",
    "R2.1": [
        "缺少理论分析。稿件没有就以下方面给出严格结果：",
        "原奇异积分方程与变换后的微分系统之间的等价性；",
        "所提出变换的稳定性；",
        "迭代神经网络过程的收敛性；",
        "在奇异点附近引入的参数 ε 的影响。",
    ],
    "R2.2": "与现有方法的比较不充分。D-AIM 的有效性仅通过基准算例得到展示，缺少与配置法、谱方法、边界元技术或现有基于 PINN 的方法等成熟数值方法的比较。",
    "R2.3": "对超参数的敏感性。网络结构、损失权重、采样策略以及 ε 的取值似乎均凭经验选定。开展敏感性研究将显著增强稿件的说服力。",
    "R2.4": "计算效率。稿件没有讨论训练时间、硬件配置、收敛历史和计算复杂度，而这些方面对于评估基于深度学习的数值方法的实际应用价值十分重要。",
    "R2.5": "文献综述应更好地说明所提出方法相对于 PINN、XPINN、BINN、DeepONet 和神经算子框架的定位。",
    "R2.6": "稿件中仍存在若干语法和文体问题，建议认真校对。",
    "R2.7": "与仅展示解的图相比，误差分布图和收敛曲线将提供更多信息。",
    "R2.8": "增加一幅说明 D-AIM 流程的示意图将有助于提高可读性。",
    "R2.9": "增加一张符号汇总表将有助于读者理解推导过程。",
}

COMMENT_BLOCKS = {
    "R1.1": [13],
    "R1.2": [18],
    "R1.3": [24],
    "R1.4": [28],
    "R1.5": [33],
    "R1.6": [37],
    "R1.7": [41],
    "R1.8": [45],
    "R1.9": [51],
    "R1.10": [56],
    "R1.11": [60],
    "R2.1": [69, 70, 71, 72, 73],
    "R2.2": [89],
    "R2.3": [96],
    "R2.4": [107],
    "R2.5": [122],
    "R2.6": [129],
    "R2.7": [136],
    "R2.8": [147],
    "R2.9": [155],
}

FORMULA_RECONSTRUCTIONS = {
    "R1.2": (
        "Page 4, line no. 16 from top, the statement “G(x,t) and φ(t) are both continuous, "
        "the integrand [G(x,t)/(x−t)^α]φ(t) is also a continuous function.” seems to be wrong. "
        "For example, if G(x,t)=(x+t)^α and φ(t)=1, how [G(x,t)/(x−t)^α]φ(t) will be a continuous function."
    ),
    "R1.3": (
        "Page 4, line no. 11, 12 from top, the domain Ω and [a, x] both are sets. So, what is purpose "
        "of writing Ω∈[a,x]? Should be used “=” or “⊆”. Similarly, for Ω∈[a,x−ε] and other such expressions."
    ),
    "R1.5": (
        "Page 3, line no. 11 from top, G(x,t)/(x−t)φ(t) should be written as [G(x,t)/(x−t)]φ(t) "
        "to avoid confusing. Similar inaccuracies are there in many places. Also do the same corrections "
        "for G(x,t)/(x−t)^α."
    ),
}


def set_fonts(run, western="Times New Roman", east_asia="Microsoft YaHei", size=11, bold=None,
              italic=None, color=BLACK):
    run.font.name = western
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), western)
    rfonts.set(qn("w:hAnsi"), western)
    rfonts.set(qn("w:cs"), western)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=110, start=142, bottom=110, end=142):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn("w:tcMar"))
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, width_dxa=TABLE_WIDTH_DXA, indent_dxa=0):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(width_dxa))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent_dxa))
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    gridcol = OxmlElement("w:gridCol")
    gridcol.set(qn("w:w"), str(width_dxa))
    grid.append(gridcol)
    for row in table.rows:
        for cell in row.cells:
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcw)
            tcw.set(qn("w:w"), str(width_dxa))
            tcw.set(qn("w:type"), "dxa")


def set_table_borders(table, color, size=6, inside=True):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    names = ["top", "left", "bottom", "right", "insideH"]
    if inside:
        names.append("insideV")
    for name in names:
        node = borders.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_no_borders(table):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:val"), "nil")
        borders.append(node)


def mark_header_row(row):
    trpr = row._tr.get_or_add_trPr()
    node = trpr.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        trpr.append(node)
    node.set(qn("w:val"), "true")


def set_paragraph_border(paragraph, side="bottom", color=BLACK, size=4):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), "2")
    edge.set(qn("w:color"), color)
    pbdr.append(edge)


def add_field(paragraph, instruction, placeholder=""):
    r1 = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    r1._r.append(begin)
    r2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    r2._r.append(instr)
    r3 = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._r.append(sep)
    if placeholder:
        set_fonts(paragraph.add_run(placeholder), size=9, color=RULE_GRAY)
    r4 = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r4._r.append(end)


def new_lower_letter_num(doc, label_text="%1."):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{(0xA1D10000 + abstract_id):08X}")
    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), f"{(0xB1D10000 + abstract_id):08X}")
    abstract.extend([nsid, tmpl])
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "lowerLetter")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), label_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "510")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "510")
    ind.set(qn("w:hanging"), "255")
    ppr.extend([tabs, ind])
    lvl.extend([start, fmt, text, suff, ppr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId")
    aid.set(qn("w:val"), str(abstract_id))
    num.append(aid)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])


def read_source(source_path: Path):
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": W}
    with zipfile.ZipFile(source_path) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    body = root.find("w:body", ns)
    blocks = [child for child in body if child.tag in {f"{{{W}}}p", f"{{{W}}}tbl"}]

    def text_of(element):
        text = "".join(t.text or "" for t in element.findall(".//w:t", ns)).strip()
        return text.replace("¦Å", "ε")

    paragraphs = [text_of(b) for b in blocks if b.tag == f"{{{W}}}p"]
    title_line = paragraphs[0]
    manuscript_line = next(x for x in paragraphs if x.startswith("Manuscript Number:"))
    title = re.sub(r"^Revisions to the paper titled\s+", "", title_line).strip()
    manuscript = manuscript_line.split(":", 1)[1].strip()
    comments = {}
    reviewer = None
    current_key = None
    in_response = False
    for text in paragraphs:
        match_reviewer = re.match(r"^Response to Reviewer #(\d+):", text)
        if match_reviewer:
            reviewer = int(match_reviewer.group(1))
            current_key = None
            in_response = False
            continue
        if reviewer is None:
            continue
        if text == "Response:":
            current_key = None
            in_response = True
            continue
        match_comment = re.match(r"^\s*(\d+)\.\s*(.*)$", text)
        if match_comment:
            current_key = f"R{reviewer}.{int(match_comment.group(1))}"
            comments[current_key] = [match_comment.group(2)]
            in_response = False
            continue
        if current_key and not in_response and re.match(r"^\s*[•·]", text):
            comments[current_key].append(re.sub(r"^\s*[•·]\s*", "", text))

    for key, replacement in FORMULA_RECONSTRUCTIONS.items():
        if key in comments:
            comments[key] = [replacement]
    expected = {f"R1.{i}" for i in range(1, 12)} | {f"R2.{i}" for i in range(1, 10)}
    if set(comments) != expected:
        raise RuntimeError(f"Reviewer-comment extraction mismatch: {sorted(comments)}")
    return manuscript, title, comments


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(4.4)
    normal.paragraph_format.line_spacing = 1.05

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(ACCENT)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h1.paragraph_format.space_before = Pt(13)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLACK)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(5)
    h2.paragraph_format.keep_with_next = True

    for name, size, left in (("TOC 1", 9.0, 0), ("TOC 2", 8.5, 0.55)):
        try:
            toc_style = doc.styles[name]
        except KeyError:
            toc_style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        toc_style.font.name = "Times New Roman"
        toc_style.font.size = Pt(size)
        toc_style.font.color.rgb = RGBColor.from_string(ACCENT)
        toc_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        toc_style.paragraph_format.left_indent = Cm(left)
        toc_style.paragraph_format.space_before = Pt(0)
        toc_style.paragraph_format.space_after = Pt(0)
        toc_style.paragraph_format.line_spacing = 1.0


def clear_container(container):
    for child in list(container._element):
        container._element.remove(child)


def add_header(section, odd, manuscript):
    header = section.header if odd else section.even_page_header
    clear_container(header)
    p = header.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.8))
    left = manuscript if odd else "审稿意见中英对照"
    right = "审稿意见中英对照" if odd else manuscript
    set_fonts(p.add_run(left), size=8.5, color=RULE_GRAY)
    p.add_run("\t")
    rr = p.add_run(right)
    set_fonts(rr, size=8.5, color=RULE_GRAY)
    set_paragraph_border(p, color=BLACK, size=4)


def add_footer(section):
    footer = section.footer
    clear_container(footer)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    add_field(p, " PAGE ", "1")
    set_fonts(p.add_run(" / "), size=9, color=BLACK)
    add_field(p, " NUMPAGES ", "1")
    for run in p.runs:
        if run.text:
            set_fonts(run, size=9, color=BLACK)


def setup_page(doc, manuscript):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.25)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(1.85)
    section.header_distance = Cm(0.95)
    section.footer_distance = Cm(1.0)
    doc.settings.odd_and_even_pages_header_footer = True
    settings = doc.settings._element
    if settings.find(qn("w:mirrorMargins")) is None:
        settings.append(OxmlElement("w:mirrorMargins"))
    add_header(section, True, manuscript)
    add_header(section, False, manuscript)
    add_footer(section)


def add_title_page(doc, manuscript, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    set_fonts(p.add_run("Bilingual Reviewer Comments"), size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_fonts(p.add_run("审稿意见中英对照"), size=18, bold=True)

    meta = doc.add_table(rows=2, cols=2)
    meta.autofit = False
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_no_borders(meta)
    labels = ["Manuscript / 稿件编号：", "Article / 稿件题目："]
    values = [manuscript, title]
    widths = [Cm(4.5), Cm(11.2)]
    for r, (label, value) in enumerate(zip(labels, values)):
        meta.rows[r].cells[0].width = widths[0]
        meta.rows[r].cells[1].width = widths[1]
        meta.rows[r].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_fonts(meta.rows[r].cells[0].paragraphs[0].add_run(label), size=10.5, bold=True)
        set_fonts(meta.rows[r].cells[1].paragraphs[0].add_run(value), size=10.5)
        for cell in meta.rows[r].cells:
            set_cell_margins(cell, 20, 50, 20, 50)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)

    note = doc.add_table(rows=1, cols=1)
    set_table_geometry(note)
    set_table_borders(note, RULE_GRAY, size=4)
    cell = note.cell(0, 0)
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, 85, 113, 85, 113)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    set_fonts(p.add_run("本文件仅汇编审稿人意见；每条先列英文原文，再列中文译文，不含作者回复与修改说明。"), size=9.5)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_fonts(p.add_run("This document contains reviewer comments only. Each English original is followed by its Chinese translation; author responses and revision notes are omitted."), size=9.5, italic=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_fonts(p.add_run("Contents"), size=15, bold=True, color=ACCENT)

    toc = doc.add_paragraph()
    toc.paragraph_format.space_after = Pt(0)
    add_field(toc, ' TOC \\o "1-2" \\h \\z \\u ', "目录将在 Word 中打开时自动更新。")
    toc.add_run().add_break(WD_BREAK.PAGE)


def add_box(doc, label, content, language):
    accent = ACCENT if language == "en" else RULE_GRAY
    fill = EN_BACK if language == "en" else ZH_BACK
    table = doc.add_table(rows=2, cols=1)
    mark_header_row(table.rows[0])
    set_table_geometry(table)
    set_table_borders(table, accent, size=5)
    table.rows[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    header_cell = table.rows[0].cells[0]
    set_cell_shading(header_cell, accent)
    set_cell_margins(header_cell, 45, 113, 45, 113)
    hp = header_cell.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.keep_with_next = True
    set_fonts(hp.add_run(label), size=9.5, bold=True, color=WHITE)

    body = table.rows[1].cells[0]
    set_cell_shading(body, fill)
    set_cell_margins(body, 90, 142, 90, 142)
    if not isinstance(content, list):
        content = [content]
    for idx, item in enumerate(content):
        p = body.paragraphs[0] if idx == 0 else body.add_paragraph()
        p.paragraph_format.space_after = Pt(2 if idx < len(content) - 1 else 0)
        p.paragraph_format.line_spacing = 1.05
        if idx > 0:
            p.style = doc.styles["List Bullet"]
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2 if idx < len(content) - 1 else 0)
        set_fonts(p.add_run(item), size=10.5)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    after.paragraph_format.line_spacing = Pt(1)


def add_comment(doc, key, english, chinese):
    number = key.split(".", 1)[1]
    heading = f"{key}: Comment {number} / 意见 {number}"
    doc.add_paragraph(heading, style="Heading 2")
    add_box(doc, "English original", english, "en")
    add_box(doc, "中文译文", chinese, "zh")


def add_reviewer(doc, reviewer_no, comments):
    doc.add_paragraph(f"Reviewer #{reviewer_no} / 审稿人 #{reviewer_no}", style="Heading 1")
    prefix = f"R{reviewer_no}."
    for key, english in comments.items():
        if key.startswith(prefix):
            add_comment(doc, key, english, TRANSLATIONS[key])


def update_fields_on_open(doc):
    settings = doc.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("source_docx", type=Path)
    ap.add_argument("output_docx", type=Path)
    args = ap.parse_args()

    manuscript, title, comments = read_source(args.source_docx)
    doc = Document()
    style_document(doc)
    setup_page(doc, manuscript)
    add_title_page(doc, manuscript, title)
    add_reviewer(doc, 1, comments)
    add_reviewer(doc, 2, comments)
    update_fields_on_open(doc)

    doc.core_properties.title = "Bilingual Reviewer Comments / 审稿意见中英对照"
    doc.core_properties.subject = f"{manuscript} reviewer comments only"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output_docx)
    print(f"Created {args.output_docx}")
    print(f"Manuscript: {manuscript}")
    print(f"Comments: {len(comments)}")


if __name__ == "__main__":
    main()
